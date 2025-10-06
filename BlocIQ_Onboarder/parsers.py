"""
BlocIQ Onboarder - File Parsers
Handles Excel, PDF, Word, and CSV parsing with intelligent content extraction
"""

import openpyxl
import pandas as pd
from docx import Document
import PyPDF2
import pdfplumber
import re
import os
from typing import Dict, List, Optional, Any
from datetime import datetime


class FileParser:
    """Base parser class for all file types"""

    def __init__(self, file_path: str):
        self.file_path = file_path
        self.file_name = os.path.basename(file_path)
        self.file_ext = os.path.splitext(file_path)[1].lower()
        self.file_size = os.path.getsize(file_path)

    def get_file_metadata(self) -> Dict[str, Any]:
        """Extract basic file metadata"""
        return {
            'file_name': self.file_name,
            'file_path': self.file_path,
            'file_type': self.file_ext,
            'file_size': self.file_size,
            'parsed_at': datetime.now().isoformat()
        }


class ExcelParser(FileParser):
    """Parse Excel files (.xlsx, .xls)"""

    def parse(self) -> Dict[str, Any]:
        """Parse Excel file and return structured data"""
        import zipfile

        # Detect if file is actually old .xls format (not a zip)
        is_old_format = False
        try:
            with open(self.file_path, 'rb') as f:
                # Check magic bytes - old Excel files start with D0 CF 11 E0 (OLE2/CFB)
                magic = f.read(4)
                if magic == b'\xD0\xCF\x11\xE0':
                    is_old_format = True
        except:
            pass

        # Try xlrd first if it's an old format file
        if is_old_format:
            try:
                df = pd.read_excel(self.file_path, sheet_name=None, engine='xlrd')

                result = {
                    **self.get_file_metadata(),
                    'sheets': list(df.keys()),
                    'total_rows': sum(len(sheet) for sheet in df.values()),
                    'data': {}
                }

                for sheet_name, sheet_df in df.items():
                    result['data'][sheet_name] = {
                        'rows': len(sheet_df),
                        'columns': list(sheet_df.columns),
                        'sample': sheet_df.head(5).to_dict('records') if not sheet_df.empty else [],
                        'raw_data': sheet_df.to_dict('records')
                    }

                return result
            except Exception as e:
                print(f"  ⚠️  xlrd failed: {e}")
                # Fall through to openpyxl
                pass

        # Try modern xlsx format
        try:
            # Try pandas first for simple tables
            df = pd.read_excel(self.file_path, sheet_name=None, engine='openpyxl')

            result = {
                **self.get_file_metadata(),
                'sheets': list(df.keys()),
                'total_rows': sum(len(sheet) for sheet in df.values()),
                'data': {}
            }

            # Parse each sheet
            for sheet_name, sheet_df in df.items():
                result['data'][sheet_name] = {
                    'rows': len(sheet_df),
                    'columns': list(sheet_df.columns),
                    'sample': sheet_df.head(5).to_dict('records') if not sheet_df.empty else [],
                    'raw_data': sheet_df.to_dict('records')
                }

            return result

        except Exception as e:
            # Final fallback to openpyxl for complex formatted sheets
            return self._parse_with_openpyxl()

    def _parse_with_openpyxl(self) -> Dict[str, Any]:
        """Parse complex Excel files with formatting"""
        try:
            wb = openpyxl.load_workbook(self.file_path, data_only=True)

            result = {
                **self.get_file_metadata(),
                'sheets': wb.sheetnames,
                'data': {}
            }

            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]

                # Extract all rows
                rows = []
                for row in ws.iter_rows(values_only=True):
                    if any(cell is not None for cell in row):
                        rows.append([str(cell) if cell is not None else '' for cell in row])

                result['data'][sheet_name] = {
                    'rows': len(rows),
                    'raw_data': rows
                }

            return result

        except Exception as e:
            return {
                **self.get_file_metadata(),
                'error': str(e),
                'data': {}
            }


class PDFParser(FileParser):
    """Parse PDF files"""

    def parse(self) -> Dict[str, Any]:
        """Parse PDF and extract text content"""
        try:
            # Try pdfplumber first (better for tables)
            return self._parse_with_pdfplumber()
        except Exception as e:
            # Fallback to PyPDF2
            return self._parse_with_pypdf2()

    def _parse_with_pdfplumber(self) -> Dict[str, Any]:
        """Parse PDF with table extraction"""
        text_content = []
        tables = []

        with pdfplumber.open(self.file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                # Extract text
                text = page.extract_text()
                if text:
                    text_content.append({
                        'page': page_num,
                        'text': text
                    })

                # Extract tables
                page_tables = page.extract_tables()
                if page_tables:
                    for table_num, table in enumerate(page_tables, 1):
                        tables.append({
                            'page': page_num,
                            'table': table_num,
                            'data': table
                        })

        return {
            **self.get_file_metadata(),
            'pages': len(text_content),
            'text_content': text_content,
            'tables': tables,
            'full_text': '\n\n'.join([p['text'] for p in text_content])
        }

    def _parse_with_pypdf2(self) -> Dict[str, Any]:
        """Fallback PDF parser"""
        text_content = []

        with open(self.file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)

            for page_num, page in enumerate(pdf_reader.pages, 1):
                text = page.extract_text()
                if text:
                    text_content.append({
                        'page': page_num,
                        'text': text
                    })

        return {
            **self.get_file_metadata(),
            'pages': len(text_content),
            'text_content': text_content,
            'full_text': '\n\n'.join([p['text'] for p in text_content])
        }


class WordParser(FileParser):
    """Parse Word documents (.docx)"""

    def parse(self) -> Dict[str, Any]:
        """Parse Word document and extract content"""
        try:
            doc = Document(self.file_path)

            # Extract paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

            # Extract tables
            tables = []
            for table_num, table in enumerate(doc.tables, 1):
                table_data = []
                for row in table.rows:
                    table_data.append([cell.text for cell in row.cells])
                tables.append({
                    'table': table_num,
                    'data': table_data
                })

            return {
                **self.get_file_metadata(),
                'paragraphs': paragraphs,
                'tables': tables,
                'full_text': '\n\n'.join(paragraphs)
            }

        except Exception as e:
            return {
                **self.get_file_metadata(),
                'error': str(e),
                'paragraphs': [],
                'tables': []
            }


class CSVParser(FileParser):
    """Parse CSV files"""

    def parse(self) -> Dict[str, Any]:
        """Parse CSV file"""
        try:
            df = pd.read_csv(self.file_path)

            return {
                **self.get_file_metadata(),
                'rows': len(df),
                'columns': list(df.columns),
                'sample': df.head(10).to_dict('records'),
                'raw_data': df.to_dict('records')
            }

        except Exception as e:
            return {
                **self.get_file_metadata(),
                'error': str(e),
                'raw_data': []
            }


def parse_file(file_path: str) -> Dict[str, Any]:
    """
    Parse any supported file type and return structured data

    Args:
        file_path: Path to the file to parse

    Returns:
        Dictionary containing parsed data and metadata
    """
    ext = os.path.splitext(file_path)[1].lower()

    parsers = {
        '.xlsx': ExcelParser,
        '.xls': ExcelParser,
        '.pdf': PDFParser,
        '.docx': WordParser,
        '.doc': WordParser,
        '.csv': CSVParser
    }

    parser_class = parsers.get(ext)

    if not parser_class:
        return {
            'file_name': os.path.basename(file_path),
            'file_path': file_path,
            'file_type': ext,
            'error': f'Unsupported file type: {ext}',
            'parsed': False
        }

    parser = parser_class(file_path)
    result = parser.parse()
    result['parsed'] = 'error' not in result

    return result


def extract_text_from_parsed(parsed_data: Dict[str, Any]) -> str:
    """Extract plain text from any parsed file data"""
    if 'full_text' in parsed_data:
        return parsed_data['full_text']

    if 'paragraphs' in parsed_data:
        return '\n\n'.join(parsed_data['paragraphs'])

    if 'data' in parsed_data and isinstance(parsed_data['data'], dict):
        # Excel file - concatenate all sheet data
        text_parts = []
        for sheet_name, sheet_data in parsed_data['data'].items():
            text_parts.append(f"Sheet: {sheet_name}")
            if 'raw_data' in sheet_data:
                text_parts.append(str(sheet_data['raw_data']))
        return '\n\n'.join(text_parts)

    return str(parsed_data)
