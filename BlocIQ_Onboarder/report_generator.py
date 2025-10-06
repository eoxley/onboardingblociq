"""
BlocIQ Onboarder - Report Generator
Exports summary.json as formatted PDF or Word document
"""

import json
import os
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, Optional


def export_summary_report(summary_path: str, output_dir: str = "output",
                          output_format: Optional[str] = None) -> str:
    """
    Generate a formatted summary report from summary.json

    Args:
        summary_path: Path to summary.json file
        output_dir: Output directory (default: "output")
        output_format: Force format ("pdf" or "word"), otherwise uses OUTPUT_FORMAT env var

    Returns:
        Path to generated report file
    """
    # Determine output format
    if output_format is None:
        output_format = os.getenv("OUTPUT_FORMAT", "pdf").lower()

    output_format = output_format.lower()

    # Load summary data
    with open(summary_path, 'r') as f:
        summary = json.load(f)

    # Ensure output directory exists
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)

    # Generate report based on format
    if output_format == "word" or output_format == "docx":
        report_file = _generate_word_report(summary, output_path)
    else:  # Default to PDF
        report_file = _generate_pdf_report(summary, output_path)

    # Update summary.json with report file reference
    summary['report_file'] = str(report_file)
    with open(summary_path, 'w') as f:
        json.dump(summary, f, indent=2)

    return str(report_file)


def _generate_pdf_report(summary: Dict[str, Any], output_path: Path) -> Path:
    """Generate PDF summary report"""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib.enums import TA_LEFT, TA_CENTER
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib import colors
    except ImportError:
        raise ImportError(
            "reportlab is required for PDF export. Install with: pip install reportlab>=4.0.0"
        )

    pdf_path = output_path / "summary_report.pdf"

    # Create PDF document
    doc = SimpleDocTemplate(
        str(pdf_path),
        pagesize=A4,
        rightMargin=2*cm,
        leftMargin=2*cm,
        topMargin=2*cm,
        bottomMargin=2*cm
    )

    # Build story (content)
    story = []
    styles = getSampleStyleSheet()

    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#1e40af'),
        spaceAfter=30,
        alignment=TA_CENTER
    )

    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=colors.HexColor('#1e40af'),
        spaceAfter=12,
        spaceBefore=12
    )

    # Title
    story.append(Paragraph("🏢 BlocIQ Building Onboarding Summary", title_style))
    story.append(Spacer(1, 0.5*cm))

    # Metadata
    building_name = summary.get('building_name', 'Unknown Building')
    timestamp = summary.get('timestamp', datetime.now().isoformat())

    story.append(Paragraph(f"<b>Building:</b> {building_name}", styles['Normal']))
    story.append(Paragraph(f"<b>Generated:</b> {timestamp}", styles['Normal']))
    story.append(Spacer(1, 0.5*cm))

    # Statistics
    story.append(Paragraph("📊 Onboarding Statistics", heading_style))

    stats = summary.get('statistics', {})
    stats_data = [
        ['Metric', 'Count'],
        ['Files Parsed', stats.get('files_parsed', 0)],
        ['Buildings', stats.get('buildings', 0)],
        ['Units', stats.get('units', 0)],
        ['Leaseholders', stats.get('leaseholders', 0)],
        ['Documents', stats.get('documents', 0)]
    ]

    stats_table = Table(stats_data, colWidths=[8*cm, 4*cm])
    stats_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1e40af')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))

    story.append(stats_table)
    story.append(Spacer(1, 0.5*cm))

    # Document Categories
    categories = summary.get('categories', {})
    if categories:
        story.append(Paragraph("📁 Document Categories", heading_style))

        cat_data = [['Category', 'Count']]
        for category, count in sorted(categories.items()):
            cat_data.append([category.replace('_', ' ').title(), count])

        cat_table = Table(cat_data, colWidths=[8*cm, 4*cm])
        cat_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1e40af')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))

        story.append(cat_table)
        story.append(Spacer(1, 0.5*cm))

    # Post-Processing Summary
    post_proc = summary.get('post_processing_summary', {})
    if post_proc:
        story.append(Paragraph("🧠 Post-Processing Enrichment", heading_style))

        enrich_data = [['Enhancement', 'Count']]
        if post_proc.get('addresses_extracted', 0) > 0:
            enrich_data.append(['Addresses Extracted', post_proc['addresses_extracted']])
        if post_proc.get('major_works_tagged', 0) > 0:
            enrich_data.append(['Major Works Tagged', post_proc['major_works_tagged']])
        if post_proc.get('budget_placeholders_created', 0) > 0:
            enrich_data.append(['Budget Placeholders', post_proc['budget_placeholders_created']])
        if post_proc.get('building_intelligence_entries', 0) > 0:
            enrich_data.append(['Building Knowledge Entries', post_proc['building_intelligence_entries']])

        if len(enrich_data) > 1:
            enrich_table = Table(enrich_data, colWidths=[8*cm, 4*cm])
            enrich_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#10b981')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.lightgreen),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))

            story.append(enrich_table)
            story.append(Spacer(1, 0.5*cm))

    # Building Intelligence
    intelligence = summary.get('building_intelligence', {})
    if intelligence and intelligence.get('total_entries', 0) > 0:
        story.append(Paragraph("🧠 Building Intelligence", heading_style))

        story.append(Paragraph(
            f"Total Entries: {intelligence['total_entries']}",
            styles['Normal']
        ))

        by_category = intelligence.get('by_category', {})
        if by_category:
            intel_data = [['Category', 'Entries']]
            for category, count in sorted(by_category.items()):
                emoji = {
                    'access': '🔑',
                    'utilities': '⚡',
                    'safety': '🚨',
                    'general': '🧾'
                }.get(category, '📋')
                intel_data.append([f"{emoji} {category.title()}", count])

            intel_table = Table(intel_data, colWidths=[8*cm, 4*cm])
            intel_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#8b5cf6')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.lavender),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))

            story.append(intel_table)
            story.append(Spacer(1, 0.5*cm))

    # Compliance Assets Breakdown
    compliance = summary.get('compliance_assets', {})
    if compliance and compliance.get('total', 0) > 0:
        story.append(Paragraph("🔍 Compliance Assets", heading_style))

        # Summary stats
        summary_data = [
            ['Metric', 'Count'],
            ['Total Assets', compliance.get('total', 0)],
            ['Compliant', compliance.get('by_status', {}).get('compliant', 0)],
            ['Overdue', compliance.get('by_status', {}).get('overdue', 0)],
            ['Due Soon', compliance.get('by_status', {}).get('due_soon', 0)],
            ['Unknown', compliance.get('by_status', {}).get('unknown', 0)]
        ]

        summary_table = Table(summary_data, colWidths=[8*cm, 4*cm])
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#dc2626')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#fee2e2')),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))

        story.append(summary_table)
        story.append(Spacer(1, 0.3*cm))

        # By Type breakdown
        by_type = compliance.get('by_type', {})
        if by_type:
            story.append(Paragraph("By Asset Type:", styles['Heading2']))
            type_data = [['Type', 'Count']]
            for asset_type, count in sorted(by_type.items(), key=lambda x: -x[1]):
                type_data.append([asset_type.replace('_', ' ').title(), count])

            type_table = Table(type_data, colWidths=[8*cm, 4*cm])
            type_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#6b7280')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f3f4f6')),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))

            story.append(type_table)
            story.append(Spacer(1, 0.5*cm))

        # Detailed asset list
        details = compliance.get('details', [])
        if details and len(details) <= 50:  # Only show detailed list if not too many
            story.append(Paragraph("Detailed Asset List:", styles['Heading2']))
            detail_data = [['Asset Name', 'Type', 'Status', 'Next Due']]
            for asset in details[:50]:  # Limit to first 50
                detail_data.append([
                    asset.get('name', 'Unknown')[:40],
                    asset.get('type', 'unknown').replace('_', ' ').title()[:20],
                    asset.get('status', 'unknown').title(),
                    asset.get('next_due', 'N/A')
                ])

            detail_table = Table(detail_data, colWidths=[6*cm, 3*cm, 2*cm, 3*cm])
            detail_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#374151')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 8),
                ('FONTSIZE', (0, 1), (-1, -1), 7),
                ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#e5e7eb')),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey)
            ]))

            story.append(detail_table)
            story.append(Spacer(1, 0.5*cm))

    # Contractor Contracts Breakdown
    contracts = summary.get('contractor_contracts', {})
    if contracts and contracts.get('total', 0) > 0:
        story.append(Paragraph("📋 Contractor Contracts", heading_style))

        # Summary stats
        contract_summary = [
            ['Metric', 'Count'],
            ['Total Contracts', contracts.get('total', 0)],
            ['Expiring Next 90 Days', contracts.get('expiring_next_90_days', 0)],
            ['Retenders Pending', contracts.get('retenders_pending', 0)],
            ['Budget Review Links', contracts.get('budget_review_links', 0)]
        ]

        contract_table = Table(contract_summary, colWidths=[8*cm, 4*cm])
        contract_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f59e0b')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#fef3c7')),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))

        story.append(contract_table)
        story.append(Spacer(1, 0.3*cm))

        # By Type breakdown
        by_type = contracts.get('by_type', {})
        if by_type:
            story.append(Paragraph("By Contractor Type:", styles['Heading2']))
            type_data = [['Type', 'Count']]
            for contractor_type, count in sorted(by_type.items(), key=lambda x: -x[1]):
                type_data.append([contractor_type.replace('_', ' ').title(), count])

            type_table = Table(type_data, colWidths=[8*cm, 4*cm])
            type_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#6b7280')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f3f4f6')),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))

            story.append(type_table)
            story.append(Spacer(1, 0.5*cm))

        # Detailed contract list
        details = contracts.get('details', [])
        if details and len(details) <= 30:  # Only show detailed list if not too many
            story.append(Paragraph("Detailed Contract List:", styles['Heading2']))
            detail_data = [['Company', 'Type', 'End Date', 'Status']]
            for contract in details[:30]:  # Limit to first 30
                detail_data.append([
                    contract.get('company', 'Unknown')[:30],
                    contract.get('type', 'unknown').replace('_', ' ').title()[:20],
                    contract.get('end_date', 'N/A'),
                    contract.get('retender_status', 'unknown').replace('_', ' ').title()
                ])

            detail_table = Table(detail_data, colWidths=[5*cm, 4*cm, 2.5*cm, 2.5*cm])
            detail_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#374151')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 8),
                ('FONTSIZE', (0, 1), (-1, -1), 7),
                ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#e5e7eb')),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey)
            ]))

            story.append(detail_table)
            story.append(Spacer(1, 0.5*cm))

    # Footer
    story.append(Spacer(1, 1*cm))
    story.append(Paragraph(
        "Generated by BlocIQ Onboarding Generator",
        ParagraphStyle('Footer', parent=styles['Normal'], alignment=TA_CENTER, textColor=colors.grey)
    ))

    # Build PDF
    doc.build(story)

    print(f"  ✅ Summary PDF created: {pdf_path}")

    return pdf_path


def _generate_word_report(summary: Dict[str, Any], output_path: Path) -> Path:
    """Generate Word document summary report"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError(
            "python-docx is required for Word export. Install with: pip install python-docx>=1.0.0"
        )

    docx_path = output_path / "summary_report.docx"

    # Create document
    doc = Document()

    # Title
    title = doc.add_heading('🏢 BlocIQ Building Onboarding Summary', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    building_name = summary.get('building_name', 'Unknown Building')
    timestamp = summary.get('timestamp', datetime.now().isoformat())

    doc.add_paragraph(f"Building: {building_name}", style='Intense Quote')
    doc.add_paragraph(f"Generated: {timestamp}", style='Intense Quote')

    # Statistics
    doc.add_heading('📊 Onboarding Statistics', level=1)

    stats = summary.get('statistics', {})
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Count'

    for metric, value in [
        ('Files Parsed', stats.get('files_parsed', 0)),
        ('Buildings', stats.get('buildings', 0)),
        ('Units', stats.get('units', 0)),
        ('Leaseholders', stats.get('leaseholders', 0)),
        ('Documents', stats.get('documents', 0))
    ]:
        row_cells = table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = str(value)

    # Document Categories
    categories = summary.get('categories', {})
    if categories:
        doc.add_heading('📁 Document Categories', level=1)

        cat_table = doc.add_table(rows=1, cols=2)
        cat_table.style = 'Light Grid Accent 1'

        hdr_cells = cat_table.rows[0].cells
        hdr_cells[0].text = 'Category'
        hdr_cells[1].text = 'Count'

        for category, count in sorted(categories.items()):
            row_cells = cat_table.add_row().cells
            row_cells[0].text = category.replace('_', ' ').title()
            row_cells[1].text = str(count)

    # Post-Processing Summary
    post_proc = summary.get('post_processing_summary', {})
    if post_proc:
        doc.add_heading('🧠 Post-Processing Enrichment', level=1)

        enrich_table = doc.add_table(rows=1, cols=2)
        enrich_table.style = 'Light Grid Accent 5'

        hdr_cells = enrich_table.rows[0].cells
        hdr_cells[0].text = 'Enhancement'
        hdr_cells[1].text = 'Count'

        if post_proc.get('addresses_extracted', 0) > 0:
            row_cells = enrich_table.add_row().cells
            row_cells[0].text = 'Addresses Extracted'
            row_cells[1].text = str(post_proc['addresses_extracted'])

        if post_proc.get('major_works_tagged', 0) > 0:
            row_cells = enrich_table.add_row().cells
            row_cells[0].text = 'Major Works Tagged'
            row_cells[1].text = str(post_proc['major_works_tagged'])

        if post_proc.get('budget_placeholders_created', 0) > 0:
            row_cells = enrich_table.add_row().cells
            row_cells[0].text = 'Budget Placeholders'
            row_cells[1].text = str(post_proc['budget_placeholders_created'])

        if post_proc.get('building_intelligence_entries', 0) > 0:
            row_cells = enrich_table.add_row().cells
            row_cells[0].text = 'Building Knowledge Entries'
            row_cells[1].text = str(post_proc['building_intelligence_entries'])

    # Building Intelligence
    intelligence = summary.get('building_intelligence', {})
    if intelligence and intelligence.get('total_entries', 0) > 0:
        doc.add_heading('🧠 Building Intelligence', level=1)

        doc.add_paragraph(f"Total Entries: {intelligence['total_entries']}")

        by_category = intelligence.get('by_category', {})
        if by_category:
            intel_table = doc.add_table(rows=1, cols=2)
            intel_table.style = 'Light Grid Accent 6'

            hdr_cells = intel_table.rows[0].cells
            hdr_cells[0].text = 'Category'
            hdr_cells[1].text = 'Entries'

            for category, count in sorted(by_category.items()):
                emoji = {
                    'access': '🔑',
                    'utilities': '⚡',
                    'safety': '🚨',
                    'general': '🧾'
                }.get(category, '📋')
                row_cells = intel_table.add_row().cells
                row_cells[0].text = f"{emoji} {category.title()}"
                row_cells[1].text = str(count)

    # Contractor Contracts Summary
    contracts = summary.get('contractor_contracts', {})
    if contracts and contracts.get('total', 0) > 0:
        doc.add_heading('📋 Contractor Contracts', level=1)

        contract_table = doc.add_table(rows=1, cols=2)
        contract_table.style = 'Light Grid Accent 3'

        hdr_cells = contract_table.rows[0].cells
        hdr_cells[0].text = 'Metric'
        hdr_cells[1].text = 'Count'

        contract_metrics = [
            ('Total Contracts', contracts.get('total', 0)),
            ('Expiring Next 90 Days', contracts.get('expiring_next_90_days', 0)),
            ('Retenders Pending', contracts.get('retenders_pending', 0)),
            ('Budget Review Links', contracts.get('budget_review_links', 0))
        ]

        for metric, count in contract_metrics:
            row_cells = contract_table.add_row().cells
            row_cells[0].text = metric
            row_cells[1].text = str(count)

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph('Generated by BlocIQ Onboarding Generator')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save document
    doc.save(str(docx_path))

    print(f"  ✅ Summary Word doc created: {docx_path}")

    return docx_path
