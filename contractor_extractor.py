#!/usr/bin/env python3
"""
Contractor Data Extractor
==========================
Extracts contractor information from Excel, PDF, or Word documents

Extracts:
- Contractor name, address, postcode
- Email, telephone
- Services provided
- Bank details (account name, sort code)
- PLI expiry date
- Document status (audited accounts, certificate of incorporation)
"""

import re
import json
from pathlib import Path
from typing import Dict, List, Optional, Any
from datetime import datetime
import PyPDF2
import openpyxl
from docx import Document


class ContractorExtractor:
    """Extract contractor information from various document formats"""
    
    def __init__(self):
        self.data = {
            'contractor_name': None,
            'address': None,
            'postcode': None,
            'email': None,
            'telephone': None,
            'services_provided': [],
            'bank_account_name': None,
            'bank_sort_code': None,
            'pli_expiry_date': None,
            'has_audited_accounts': False,
            'has_certificate_of_incorporation': False,
            'source_document': None,
            'extraction_method': None,
            'extraction_confidence': 0.0,
            'documents_found': []
        }
        
        # Common service types
        self.service_keywords = [
            'cleaning', 'lift maintenance', 'electrical', 'plumbing', 'heating',
            'cctv', 'security', 'gardening', 'grounds maintenance', 'pest control',
            'fire alarm', 'fire safety', 'gas safety', 'water hygiene', 'drainage',
            'roofing', 'painting', 'decorating', 'carpentry', 'joinery'
        ]
    
    def extract_from_files(self, files: List[Path]) -> Dict:
        """
        Extract contractor data from a list of files
        
        Args:
            files: List of Path objects pointing to files
            
        Returns:
            Dictionary with extracted contractor data
        """
        print(f"\n📁 Processing {len(files)} file(s)")
        
        # Extract from each file
        for file in files:
            print(f"\n   📄 Processing: {file.name}")
            ext_lower = file.suffix.lower()
            
            try:
                if ext_lower in ['.xlsx', '.xls']:
                    print(f"      → Excel file detected")
                    self._extract_from_excel(file)
                elif ext_lower == '.pdf':
                    print(f"      → PDF file detected")
                    self._extract_from_pdf(file)
                elif ext_lower in ['.docx', '.doc']:
                    print(f"      → Word file detected")
                    self._extract_from_word(file)
                else:
                    print(f"      ⚠️  Unsupported file type: {ext_lower}")
                    continue
            except Exception as e:
                print(f"      ⚠️  Error processing file: {e}")
            
            # Track documents
            self.data['documents_found'].append({
                'file_name': file.name,
                'file_type': file.suffix.lower(),
                'file_size': file.stat().st_size
            })
        
        # Post-processing
        self._infer_document_status(files)
        self._calculate_confidence()
        
        return self.data
    
    def extract_from_folder(self, folder_path: str) -> Dict:
        """
        Extract contractor data from a folder containing contractor documents
        
        Args:
            folder_path: Path to folder with contractor documents
            
        Returns:
            Dictionary with extracted contractor data
        """
        folder = Path(folder_path)
        if not folder.exists():
            raise FileNotFoundError(f"Folder not found: {folder_path}")
        
        print(f"\n📁 Scanning folder: {folder.name}")
        
        # Find all relevant files (case-insensitive)
        files = []
        for file in folder.iterdir():
            if file.is_file():
                ext_lower = file.suffix.lower()
                if ext_lower in ['.xlsx', '.xls', '.pdf', '.docx', '.doc']:
                    files.append(file)
        
        print(f"   Found {len(files)} document(s)")
        
        if not files:
            print("   ⚠️  No Excel, PDF, or Word files found in folder")
            return self.data
        
        # Extract from each file
        for file in files:
            print(f"\n   📄 Processing: {file.name}")
            ext_lower = file.suffix.lower()
            
            try:
                if ext_lower in ['.xlsx', '.xls']:
                    print(f"      → Excel file detected")
                    self._extract_from_excel(file)
                elif ext_lower == '.pdf':
                    print(f"      → PDF file detected")
                    self._extract_from_pdf(file)
                elif ext_lower in ['.docx', '.doc']:
                    print(f"      → Word file detected")
                    self._extract_from_word(file)
            except Exception as e:
                print(f"      ⚠️  Error processing file: {e}")
            
            # Track documents
            self.data['documents_found'].append({
                'file_name': file.name,
                'file_type': file.suffix.lower(),
                'file_size': file.stat().st_size
            })
        
        # Post-processing
        self._infer_document_status(files)
        self._calculate_confidence()
        
        return self.data
    
    def _extract_from_excel(self, file_path: Path):
        """Extract data from Excel file"""
        try:
            print(f"      → Opening Excel workbook...")
            
            # Try with data_only first, then without
            try:
                wb = openpyxl.load_workbook(file_path, data_only=True)
            except:
                wb = openpyxl.load_workbook(file_path)
            
            ws = wb.active
            print(f"      → Reading worksheet: {ws.title}")
            
            # Read all cell values
            text_content = []
            rows_read = 0
            
            for row in ws.iter_rows(max_row=200):  # Increased to 200 rows
                row_text = []
                for cell in row:
                    if cell.value is not None:
                        row_text.append(str(cell.value))
                
                if row_text:
                    text_content.append(' | '.join(row_text))
                    rows_read += 1
            
            print(f"      → Read {rows_read} rows")
            
            # Join and extract
            full_text = '\n'.join(text_content)
            
            if len(full_text) < 50:
                print(f"      ⚠️  Very little text extracted ({len(full_text)} chars)")
            else:
                print(f"      → Extracted {len(full_text)} characters")
            
            self._extract_fields_from_text(full_text, 'excel')
            
            self.data['extraction_method'] = 'excel'
            print(f"      ✓ Excel data extracted successfully")
            
            wb.close()
            
        except PermissionError:
            print(f"      ⚠️  Permission denied - file may be open in Excel")
        except Exception as e:
            print(f"      ⚠️  Could not read Excel: {e}")
            import traceback
            traceback.print_exc()
    
    def _extract_from_pdf(self, file_path: Path):
        """Extract data from PDF file"""
        try:
            print(f"      → Opening PDF...")
            
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                num_pages = len(reader.pages)
                print(f"      → PDF has {num_pages} pages")
                
                text_content = []
                
                for i, page in enumerate(reader.pages[:10], 1):  # First 10 pages
                    text = page.extract_text()
                    if text:
                        text_content.append(text)
                        print(f"      → Page {i}: {len(text)} chars")
                
                full_text = '\n'.join(text_content)
                
                if len(full_text) < 50:
                    print(f"      ⚠️  Very little text extracted ({len(full_text)} chars)")
                else:
                    print(f"      → Total extracted: {len(full_text)} characters")
                
                self._extract_fields_from_text(full_text, 'pdf_ocr')
            
            self.data['extraction_method'] = 'pdf_ocr'
            print(f"      ✓ PDF data extracted successfully")
            
        except Exception as e:
            print(f"      ⚠️  Could not read PDF: {e}")
            import traceback
            traceback.print_exc()
    
    def _extract_from_word(self, file_path: Path):
        """Extract data from Word document"""
        try:
            print(f"      → Opening Word document...")
            
            doc = Document(file_path)
            text_content = []
            
            # Read paragraphs
            print(f"      → Reading {len(doc.paragraphs)} paragraphs")
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
            
            # Also check tables
            if doc.tables:
                print(f"      → Reading {len(doc.tables)} tables")
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text_content.append(cell.text)
            
            full_text = '\n'.join(text_content)
            
            if len(full_text) < 50:
                print(f"      ⚠️  Very little text extracted ({len(full_text)} chars)")
            else:
                print(f"      → Total extracted: {len(full_text)} characters")
            
            self._extract_fields_from_text(full_text, 'word')
            
            self.data['extraction_method'] = 'word'
            print(f"      ✓ Word data extracted successfully")
            
        except Exception as e:
            print(f"      ⚠️  Could not read Word document: {e}")
            import traceback
            traceback.print_exc()
    
    def _extract_fields_from_text(self, text: str, method: str):
        """Extract specific fields from text using regex patterns"""
        
        # Contractor Name (often appears as company name at top)
        if not self.data['contractor_name']:
            # Look for "Limited", "Ltd", "LLP", etc.
            company_pattern = r'([A-Z][A-Za-z\s&]+(?:Limited|Ltd|LLP|plc|PLC))'
            match = re.search(company_pattern, text)
            if match:
                self.data['contractor_name'] = match.group(1).strip()
        
        # Email
        if not self.data['email']:
            email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
            match = re.search(email_pattern, text)
            if match:
                self.data['email'] = match.group(0)
        
        # Telephone (UK format)
        if not self.data['telephone']:
            phone_patterns = [
                r'\b0\d{10}\b',  # 01234567890
                r'\b0\d{3}\s?\d{3}\s?\d{4}\b',  # 0123 456 7890
                r'\+44\s?\d{3,4}\s?\d{3,4}\s?\d{4}\b'  # +44 123 456 7890
            ]
            for pattern in phone_patterns:
                match = re.search(pattern, text)
                if match:
                    self.data['telephone'] = match.group(0)
                    break
        
        # Postcode (UK format)
        if not self.data['postcode']:
            postcode_pattern = r'\b[A-Z]{1,2}\d{1,2}[A-Z]?\s?\d[A-Z]{2}\b'
            match = re.search(postcode_pattern, text)
            if match:
                self.data['postcode'] = match.group(0)
        
        # Sort Code (XX-XX-XX format)
        if not self.data['bank_sort_code']:
            sort_code_pattern = r'\b\d{2}[-\s]?\d{2}[-\s]?\d{2}\b'
            matches = re.findall(sort_code_pattern, text)
            if matches:
                # Look for one near "sort code" text
                for i, line in enumerate(text.split('\n')):
                    if 'sort code' in line.lower():
                        match = re.search(sort_code_pattern, line)
                        if match:
                            self.data['bank_sort_code'] = match.group(0)
                            break
                if not self.data['bank_sort_code'] and matches:
                    self.data['bank_sort_code'] = matches[0]
        
        # PLI Expiry Date
        if not self.data['pli_expiry_date']:
            # Look for dates near "expiry", "expires", "renewal"
            date_patterns = [
                r'\d{1,2}[/-]\d{1,2}[/-]\d{2,4}',  # DD/MM/YYYY or DD-MM-YYYY
                r'\d{4}[/-]\d{1,2}[/-]\d{1,2}',    # YYYY-MM-DD
            ]
            
            for line in text.split('\n'):
                if any(word in line.lower() for word in ['expiry', 'expires', 'expiration', 'renewal', 'valid until']):
                    for pattern in date_patterns:
                        match = re.search(pattern, line)
                        if match:
                            self.data['pli_expiry_date'] = match.group(0)
                            break
                    if self.data['pli_expiry_date']:
                        break
        
        # Services Provided
        for keyword in self.service_keywords:
            if keyword in text.lower() and keyword not in [s.lower() for s in self.data['services_provided']]:
                self.data['services_provided'].append(keyword.title())
        
        # Bank Account Name (often near "account name" or "account holder")
        if not self.data['bank_account_name']:
            for line in text.split('\n'):
                if any(phrase in line.lower() for phrase in ['account name', 'account holder', 'beneficiary']):
                    # Next line or same line after colon
                    if ':' in line:
                        parts = line.split(':')
                        if len(parts) > 1:
                            name = parts[1].strip()
                            if len(name) > 3 and not name.isdigit():
                                self.data['bank_account_name'] = name
    
    def _infer_document_status(self, files: List[Path]):
        """Infer what documents are present based on file names"""
        
        for file in files:
            name_lower = file.name.lower()
            
            # Check for audited accounts
            if any(word in name_lower for word in ['accounts', 'financial statement', 'annual report', 'audit']):
                self.data['has_audited_accounts'] = True
            
            # Check for certificate of incorporation
            if any(word in name_lower for word in ['incorporation', 'companies house', 'certificate']):
                self.data['has_certificate_of_incorporation'] = True
    
    def _calculate_confidence(self):
        """Calculate extraction confidence score"""
        
        required_fields = [
            'contractor_name',
            'email',
            'telephone',
            'postcode'
        ]
        
        fields_found = sum(1 for field in required_fields if self.data.get(field))
        confidence = fields_found / len(required_fields)
        
        # Bonus for additional fields
        if self.data.get('services_provided'):
            confidence += 0.1
        if self.data.get('bank_sort_code'):
            confidence += 0.1
        if self.data.get('pli_expiry_date'):
            confidence += 0.1
        
        self.data['extraction_confidence'] = min(1.0, confidence)


def main():
    """CLI entry point"""
    import sys
    
    if len(sys.argv) != 2:
        print("Usage: python3 contractor_extractor.py <contractor_folder>")
        print("\nExample:")
        print('  python3 contractor_extractor.py "/path/to/contractor/documents"')
        sys.exit(1)
    
    folder_path = sys.argv[1]
    
    # Extract
    extractor = ContractorExtractor()
    data = extractor.extract_from_folder(folder_path)
    
    # Print results
    print("\n" + "="*80)
    print("📊 EXTRACTION RESULTS")
    print("="*80)
    
    print(f"\n✅ Contractor Name: {data.get('contractor_name') or 'Not found'}")
    print(f"✅ Email: {data.get('email') or 'Not found'}")
    print(f"✅ Telephone: {data.get('telephone') or 'Not found'}")
    print(f"✅ Postcode: {data.get('postcode') or 'Not found'}")
    print(f"✅ Services: {', '.join(data.get('services_provided', [])) or 'Not found'}")
    print(f"✅ Bank Sort Code: {data.get('bank_sort_code') or 'Not found'}")
    print(f"✅ Bank Account Name: {data.get('bank_account_name') or 'Not found'}")
    print(f"✅ PLI Expiry: {data.get('pli_expiry_date') or 'Not found'}")
    print(f"✅ Has Audited Accounts: {data.get('has_audited_accounts')}")
    print(f"✅ Has Certificate of Incorporation: {data.get('has_certificate_of_incorporation')}")
    
    print(f"\n📊 Confidence: {data.get('extraction_confidence', 0):.0%}")
    print(f"📁 Documents: {len(data.get('documents_found', []))}")
    
    # Save to JSON
    output_file = Path(folder_path).parent / f"{Path(folder_path).name}_contractor_data.json"
    with open(output_file, 'w') as f:
        json.dump(data, f, indent=2)
    
    print(f"\n✅ Data saved to: {output_file}")
    
    return data


if __name__ == '__main__':
    main()

