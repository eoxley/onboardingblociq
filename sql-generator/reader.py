"""
SQL Generator - File Reading Layer
Modular file reader supporting PDF, Word, Excel, CSV, TXT, RTF, images (OCR), and emails
"""

import os
import mimetypes
from pathlib import Path
from typing import Dict, Any, Optional
import pdfplumber
import PyPDF2
from docx import Document
import openpyxl
import pandas as pd
import email
from email import policy
from email.parser import BytesParser
import base64
import requests
import json


class FileReader:
    """Universal file reader dispatcher"""

    SUPPORTED_EXTENSIONS = {
        '.pdf', '.doc', '.docx', '.xls', '.xlsx', '.csv',
        '.txt', '.rtf', '.msg', '.eml', '.png', '.jpg', '.jpeg'
    }

    def __init__(self, ocr_enabled: bool = True):
        """
        Initialize file reader

        Args:
            ocr_enabled: Whether to enable OCR for image files
        """
        self.ocr_enabled = ocr_enabled
        self.ocr_token = os.getenv('RENDER_OCR_TOKEN', '1')

    def read_file(self, file_path: str) -> Dict[str, Any]:
        """
        Read any supported file type and extract text

        Args:
            file_path: Path to the file

        Returns:
            Dictionary with extracted text and metadata
        """
        path = Path(file_path)

        if not path.exists():
            return {
                'file_path': file_path,
                'file_name': path.name,
                'error': 'File not found',
                'text': ''
            }

        ext = path.suffix.lower()
        size = path.stat().st_size

        # Basic metadata
        result = {
            'file_path': str(path),
            'file_name': path.name,
            'extension': ext,
            'size': size,
            'mime_type': mimetypes.guess_type(file_path)[0]
        }

        # Dispatch to appropriate reader
        try:
            if ext == '.pdf':
                text = self._read_pdf(file_path)
            elif ext in ['.docx', '.doc']:
                text = self._read_doc(file_path)
            elif ext in ['.xlsx', '.xls']:
                text = self._read_excel(file_path)
            elif ext == '.csv':
                text = self._read_csv(file_path)
            elif ext in ['.txt', '.rtf']:
                text = self._read_text(file_path)
            elif ext in ['.msg', '.eml']:
                text = self._read_email(file_path)
            elif ext in ['.png', '.jpg', '.jpeg']:
                text = self._read_image_ocr(file_path)
            else:
                text = ''
                result['error'] = f'Unsupported file type: {ext}'

            result['text'] = text
            result['success'] = True

        except Exception as e:
            result['error'] = str(e)
            result['text'] = ''
            result['success'] = False

        return result

    def _read_pdf(self, file_path: str) -> str:
        """Extract text from PDF using pdfplumber"""
        text_parts = []

        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
        except Exception as e:
            # Fallback to PyPDF2
            try:
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
            except:
                raise e

        return '\n\n'.join(text_parts)

    def _read_doc(self, file_path: str) -> str:
        """Extract text from Word document"""
        doc = Document(file_path)

        text_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text for cell in row.cells)
                text_parts.append(row_text)

        return '\n'.join(text_parts)

    def _read_excel(self, file_path: str) -> str:
        """Extract text from Excel file"""
        text_parts = []

        try:
            # Try with pandas first
            sheets = pd.read_excel(file_path, sheet_name=None, engine='openpyxl')

            for sheet_name, df in sheets.items():
                text_parts.append(f"Sheet: {sheet_name}")
                # Convert dataframe to text
                text_parts.append(df.to_string(index=False))

        except Exception as e:
            # Fallback to openpyxl
            wb = openpyxl.load_workbook(file_path, data_only=True)

            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                text_parts.append(f"Sheet: {sheet_name}")

                for row in ws.iter_rows(values_only=True):
                    if any(cell is not None for cell in row):
                        row_text = ' | '.join(str(cell) if cell is not None else '' for cell in row)
                        text_parts.append(row_text)

        return '\n'.join(text_parts)

    def _read_csv(self, file_path: str) -> str:
        """Extract text from CSV file"""
        df = pd.read_csv(file_path)
        return df.to_string(index=False)

    def _read_text(self, file_path: str) -> str:
        """Read plain text or RTF file"""
        # Try different encodings
        encodings = ['utf-8', 'latin-1', 'cp1252']

        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue

        # If all fail, read as binary and decode with errors='ignore'
        with open(file_path, 'rb') as f:
            return f.read().decode('utf-8', errors='ignore')

    def _read_email(self, file_path: str) -> str:
        """Extract text from email file (.msg or .eml)"""
        ext = Path(file_path).suffix.lower()

        if ext == '.eml':
            # Parse .eml file
            with open(file_path, 'rb') as f:
                msg = BytesParser(policy=policy.default).parse(f)

            text_parts = []

            # Subject
            if msg['subject']:
                text_parts.append(f"Subject: {msg['subject']}")

            # From
            if msg['from']:
                text_parts.append(f"From: {msg['from']}")

            # Date
            if msg['date']:
                text_parts.append(f"Date: {msg['date']}")

            # Body
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == 'text/plain':
                        text_parts.append(part.get_content())
            else:
                text_parts.append(msg.get_content())

            return '\n\n'.join(text_parts)

        elif ext == '.msg':
            # For .msg files, try extract_msg library
            try:
                import extract_msg
                msg = extract_msg.Message(file_path)

                text_parts = []
                text_parts.append(f"Subject: {msg.subject}")
                text_parts.append(f"From: {msg.sender}")
                text_parts.append(f"Date: {msg.date}")
                text_parts.append(msg.body)

                return '\n\n'.join(text_parts)
            except ImportError:
                # If extract_msg not available, return error
                return f"Error: extract_msg library required for .msg files"
            except Exception as e:
                return f"Error reading .msg file: {str(e)}"

        return ""

    def _read_image_ocr(self, file_path: str) -> str:
        """Extract text from image using OCR"""
        if not self.ocr_enabled:
            return ""

        # Use Anthropic Vision API or similar OCR service
        try:
            # Read image as base64
            with open(file_path, 'rb') as f:
                image_data = base64.standard_b64encode(f.read()).decode('utf-8')

            # Determine media type
            ext = Path(file_path).suffix.lower()
            media_types = {
                '.png': 'image/png',
                '.jpg': 'image/jpeg',
                '.jpeg': 'image/jpeg'
            }
            media_type = media_types.get(ext, 'image/jpeg')

            # For now, return placeholder - actual OCR implementation depends on service
            # In production, this would call Anthropic's Vision API or similar
            return f"[OCR extraction from {Path(file_path).name} - requires Vision API integration]"

        except Exception as e:
            return f"Error performing OCR: {str(e)}"


def scan_folder(folder_path: str, recursive: bool = True) -> list:
    """
    Recursively scan folder for all supported files

    Args:
        folder_path: Path to folder to scan
        recursive: Whether to scan subfolders

    Returns:
        List of file paths
    """
    folder = Path(folder_path)
    files = []

    pattern = '**/*' if recursive else '*'

    for item in folder.glob(pattern):
        # Skip hidden and system files
        if item.name.startswith('.') or item.name.startswith('~'):
            continue

        # Check if file has supported extension
        if item.is_file() and item.suffix.lower() in FileReader.SUPPORTED_EXTENSIONS:
            files.append(str(item))

    return sorted(files)


# Main function for testing
if __name__ == '__main__':
    import sys

    if len(sys.argv) < 2:
        print("Usage: python reader.py <file_or_folder_path>")
        sys.exit(1)

    path = sys.argv[1]
    reader = FileReader()

    if os.path.isdir(path):
        print(f"Scanning folder: {path}")
        files = scan_folder(path)
        print(f"Found {len(files)} files")

        for file_path in files[:5]:  # Show first 5
            result = reader.read_file(file_path)
            print(f"\n{result['file_name']}: {len(result.get('text', ''))} characters")
    else:
        result = reader.read_file(path)
        print(f"File: {result['file_name']}")
        print(f"Text length: {len(result.get('text', ''))}")
        print(f"Preview: {result.get('text', '')[:500]}")
